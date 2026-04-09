"""File generation service — HTML-to-PDF and DOCX generation."""
import os
import uuid
import logging
from pathlib import Path
from typing import Optional
from jinja2 import Environment, FileSystemLoader
from backend.config import settings

logger = logging.getLogger(__name__)

jinja_env = Environment(
    loader=FileSystemLoader(settings.TEMPLATE_DIR),
    autoescape=False,  # Disable autoescape for resume HTML rendering
)

# Max file sizes (in bytes)
MAX_PDF_SIZE = 5 * 1024 * 1024  # 5MB
MAX_DOCX_SIZE = 10 * 1024 * 1024  # 10MB
MAX_TXT_SIZE = 2 * 1024 * 1024  # 2MB


class FileService:

    def render_template(self, template_name: str, context: dict) -> str:
        """Render an HTML template with context data."""
        try:
            # Normalize skills format: support both dict-of-lists and dict-of-strings
            if "skills" in context and isinstance(context["skills"], dict):
                normalized = {}
                for cat, items in context["skills"].items():
                    if isinstance(items, list):
                        normalized[cat] = ", ".join(items)
                    else:
                        normalized[cat] = str(items)
                context["skills"] = normalized

            template = jinja_env.get_template(template_name)
            return template.render(**context)
        except Exception as e:
            logger.error(f"Template rendering error: {e}")
            return f"<html><body><p>Template error: {e}</p></body></html>"

    def _validate_file_size(self, filepath: str, max_size: int, file_type: str) -> bool:
        """Check if file exceeds max size and log appropriately."""
        try:
            if not os.path.exists(filepath):
                logger.error(f"File not found: {filepath}")
                return False
            
            file_size = os.path.getsize(filepath)
            if file_size > max_size:
                logger.warning(
                    f"{file_type} generation warning: File size {file_size} bytes exceeds limit {max_size} bytes",
                    extra={
                        "file_path": filepath,
                        "file_size": file_size,
                        "max_size": max_size,
                        "file_type": file_type
                    }
                )
                return False
            
            logger.debug(f"{file_type} file validated: {file_size} bytes")
            return True
        except Exception as e:
            logger.error(f"Size validation error: {e}")
            return False

    def generate_pdf(self, html_content: str, filename: Optional[str] = None) -> str:
        """Generate PDF from HTML content. Returns file path."""
        if not filename:
            filename = f"resume_{uuid.uuid4().hex[:8]}.pdf"
        filepath = os.path.join(settings.GENERATED_DIR, filename)

        try:
            from weasyprint import HTML
            HTML(string=html_content).write_pdf(filepath)
            
            # Validate file size
            if not self._validate_file_size(filepath, MAX_PDF_SIZE, "PDF"):
                os.remove(filepath)
                raise ValueError(f"Generated PDF exceeds {MAX_PDF_SIZE} bytes limit")
            
            logger.info(f"PDF generated successfully: {filepath}")
            return filepath
        except ImportError:
            logger.warning("weasyprint not installed. Trying pdfkit fallback...")
            # Try pdfkit as second fallback
            try:
                import pdfkit
                pdfkit.from_string(html_content, filepath)
                
                # Validate file size
                if not self._validate_file_size(filepath, MAX_PDF_SIZE, "PDF"):
                    os.remove(filepath)
                    raise ValueError(f"Generated PDF exceeds {MAX_PDF_SIZE} bytes limit")
                
                logger.info(f"PDF generated via pdfkit: {filepath}")
                return filepath
            except (ImportError, Exception) as e2:
                logger.warning(f"pdfkit also failed: {e2}. Saving as HTML fallback.")
                html_path = filepath.replace(".pdf", ".html")
                with open(html_path, "w", encoding="utf-8") as f:
                    f.write(html_content)
                logger.info(f"HTML fallback generated: {html_path}")
                return html_path
        except Exception as e:
            logger.error(f"PDF generation error: {e}")
            # Fallback: save HTML
            html_path = filepath.replace(".pdf", ".html")
            with open(html_path, "w", encoding="utf-8") as f:
                f.write(html_content)
            logger.info(f"HTML fallback generated due to error: {html_path}")
            return html_path

    def generate_docx(self, resume_data: dict, filename: Optional[str] = None) -> str:
        """Generate DOCX from resume data. Returns file path."""
        if not filename:
            filename = f"resume_{uuid.uuid4().hex[:8]}.docx"
        filepath = os.path.join(settings.GENERATED_DIR, filename)

        try:
            from docx import Document
            from docx.shared import Pt, Inches, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            doc = Document()

            # Style setup
            style = doc.styles['Normal']
            style.font.name = 'Calibri'
            style.font.size = Pt(11)

            # Name
            name = resume_data.get("full_name", "Your Name")
            heading = doc.add_heading(name, level=0)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Contact
            contact = resume_data.get("contact", {})
            contact_parts = []
            for key in ["email", "phone", "linkedin", "github", "portfolio", "location", "leetcode"]:
                val = contact.get(key)
                if val:
                    contact_parts.append(val)
            if contact_parts:
                p = doc.add_paragraph(" | ".join(contact_parts))
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Summary
            if resume_data.get("summary"):
                doc.add_heading("Professional Summary", level=1)
                summary = resume_data["summary"]
                if isinstance(summary, list):
                    for s in summary:
                        doc.add_paragraph(s)
                else:
                    doc.add_paragraph(str(summary))

            # Education
            education = resume_data.get("education", [])
            if education:
                doc.add_heading("Education", level=1)
                for edu in education:
                    p = doc.add_paragraph()
                    run = p.add_run(f"{edu.get('school', '')} — {edu.get('degree', '')}")
                    run.bold = True
                    details = []
                    if edu.get("dates"):
                        details.append(edu["dates"])
                    if edu.get("grade"):
                        details.append(edu["grade"])
                    if details:
                        p.add_run(f"\n{' | '.join(details)}")

            # Technical Skills
            skills = resume_data.get("skills", {})
            if skills:
                doc.add_heading("Technical Skills", level=1)
                if isinstance(skills, dict):
                    for category, skill_list in skills.items():
                        if skill_list:
                            if isinstance(skill_list, list):
                                skill_str = ", ".join(skill_list)
                            else:
                                skill_str = str(skill_list)
                            p = doc.add_paragraph()
                            run = p.add_run(f"{category}: ")
                            run.bold = True
                            p.add_run(skill_str)
                elif isinstance(skills, list):
                    doc.add_paragraph(", ".join(skills))
                else:
                    doc.add_paragraph(str(skills))

            # Experience
            experience = resume_data.get("experience", [])
            if experience:
                doc.add_heading("Experience", level=1)
                for exp in experience:
                    p = doc.add_paragraph()
                    run = p.add_run(f"{exp.get('company', '')} — {exp.get('title', '')}")
                    run.bold = True
                    sub_parts = []
                    if exp.get("location"):
                        sub_parts.append(exp["location"])
                    if exp.get("dates"):
                        sub_parts.append(exp["dates"])
                    if sub_parts:
                        p.add_run(f"\n{' | '.join(sub_parts)}")
                    for bullet in exp.get("bullets", []):
                        doc.add_paragraph(bullet, style='List Bullet')

            # Projects
            projects = resume_data.get("projects", [])
            if projects:
                doc.add_heading("Projects", level=1)
                for proj in projects:
                    p = doc.add_paragraph()
                    run = p.add_run(proj.get("name", ""))
                    run.bold = True
                    tech = proj.get("tech_stack", "")
                    if tech:
                        p.add_run(f" | {tech}")
                    
                    link_line = []
                    if proj.get("live_url"):
                        link_line.append(f"Live: {proj['live_url']}")
                    if proj.get("repo_url"):
                        link_line.append(f"Repo: {proj['repo_url']}")
                    if link_line:
                        doc.add_paragraph(" | ".join(link_line))
                        
                    for bullet in proj.get("bullets", []):
                        doc.add_paragraph(bullet, style='List Bullet')

            # Certifications
            certs = resume_data.get("certifications", [])
            if certs:
                doc.add_heading("Certifications", level=1)
                for cert in certs:
                    parts = [cert.get("name", "")]
                    if cert.get("issuer"):
                        parts.append(cert["issuer"])
                    if cert.get("date"):
                        parts.append(cert["date"])
                    doc.add_paragraph(" — ".join(parts), style='List Bullet')

            # Achievements
            achievements = resume_data.get("achievements", [])
            if achievements:
                doc.add_heading("Achievements", level=1)
                for ach in achievements:
                    doc.add_paragraph(ach, style='List Bullet')

            doc.save(filepath)
            
            # Validate file size
            if not self._validate_file_size(filepath, MAX_DOCX_SIZE, "DOCX"):
                os.remove(filepath)
                raise ValueError(f"Generated DOCX exceeds {MAX_DOCX_SIZE} bytes limit")
            
            logger.info(f"DOCX generated successfully: {filepath}")
            return filepath

        except ImportError:
            logger.error("python-docx not installed")
            return ""
        except Exception as e:
            logger.error(f"DOCX generation error: {e}")
            return ""

    def generate_txt(self, resume_data: dict, filename: Optional[str] = None) -> str:
        """Generate plain text resume."""
        if not filename:
            filename = f"resume_{uuid.uuid4().hex[:8]}.txt"
        filepath = os.path.join(settings.GENERATED_DIR, filename)

        try:
            lines = []
            lines.append(resume_data.get("full_name", "Your Name").upper())
            contact = resume_data.get("contact", {})
            contact_parts = [v for v in contact.values() if v]
            lines.append(" | ".join(contact_parts))
            lines.append("=" * 60)

            if resume_data.get("summary"):
                lines.append("\nPROFESSIONAL SUMMARY")
                lines.append("-" * 40)
                summary = resume_data["summary"]
                if isinstance(summary, list):
                    for s in summary:
                        lines.append(s)
                else:
                    lines.append(str(summary))

            if resume_data.get("education"):
                lines.append("\nEDUCATION")
                lines.append("-" * 40)
                for edu in resume_data.get("education", []):
                    lines.append(f"{edu.get('degree', '')} — {edu.get('school', '')} ({edu.get('dates', '')})")
                    if edu.get("grade"):
                        lines.append(f"  {edu['grade']}")

            skills = resume_data.get("skills", {})
            if skills:
                lines.append("\nTECHNICAL SKILLS")
                lines.append("-" * 40)
                if isinstance(skills, dict):
                    for cat, slist in skills.items():
                        if slist:
                            if isinstance(slist, list):
                                lines.append(f"{cat}: {', '.join(slist)}")
                            else:
                                lines.append(f"{cat}: {slist}")

            for exp in resume_data.get("experience", []):
                lines.append(f"\n{exp.get('company', '')} — {exp.get('title', '')}")
                lines.append(f"{exp.get('location', '')} | {exp.get('dates', '')}")
                for bullet in exp.get("bullets", []):
                    lines.append(f"  • {bullet}")

            if resume_data.get("projects"):
                lines.append("\nPROJECTS")
                lines.append("-" * 40)
                for proj in resume_data.get("projects", []):
                    tech = f" ({proj.get('tech_stack', '')})" if proj.get('tech_stack') else ""
                    lines.append(f"{proj.get('name', '')}{tech}")
                    if proj.get("live_url"):
                        lines.append(f"  Live: {proj['live_url']}")
                    for bullet in proj.get("bullets", []):
                        lines.append(f"  • {bullet}")

            if resume_data.get("certifications"):
                lines.append("\nCERTIFICATIONS")
                lines.append("-" * 40)
                for cert in resume_data.get("certifications", []):
                    parts = [cert.get("name", "")]
                    if cert.get("issuer"):
                        parts.append(cert["issuer"])
                    lines.append(" — ".join(parts))

            if resume_data.get("achievements"):
                lines.append("\nACHIEVEMENTS")
                lines.append("-" * 40)
                for ach in resume_data.get("achievements", []):
                    lines.append(f"  • {ach}")

            with open(filepath, "w", encoding="utf-8") as f:
                f.write("\n".join(lines))
            
            # Validate file size
            if not self._validate_file_size(filepath, MAX_TXT_SIZE, "TXT"):
                os.remove(filepath)
                raise ValueError(f"Generated TXT exceeds {MAX_TXT_SIZE} bytes limit")
            
            logger.info(f"TXT generated successfully: {filepath}")
            return filepath
        
        except Exception as e:
            logger.error(f"TXT generation error: {e}")
            return ""


file_service = FileService()
